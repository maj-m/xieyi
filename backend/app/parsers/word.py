"""DOCX 解析器：提取段落和表格文本；旧版二进制 DOC 暂不在进程内解析。"""

from pathlib import Path

from docx import Document  # type: ignore[import-untyped]

from app.errors import EvidenceValidationError
from app.models.evidence import Evidence
from app.parsers.base import ParseResult
from app.parsers.limits import MAX_SPREADSHEET_CELLS, MAX_WORD_PARAGRAPHS, validate_zip_container
from app.parsers.text import ensure_text_limit


class DocxParser:
    name = "python-docx"
    version = "1.0.0"

    def supports(self, evidence: Evidence) -> bool:
        return evidence.file_extension == ".docx"

    def parse(self, source: Path, evidence: Evidence) -> ParseResult:
        validate_zip_container(source)
        document = Document(str(source))
        if len(document.paragraphs) > MAX_WORD_PARAGRAPHS:
            raise EvidenceValidationError(
                "WORD_PARAGRAPH_LIMIT", "Word document has too many paragraphs"
            )
        lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells = 0
        for table_index, table in enumerate(document.tables, start=1):
            lines.append(f"## Table: {table_index}")
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                table_cells += len(values)
                if table_cells > MAX_SPREADSHEET_CELLS:
                    raise EvidenceValidationError(
                        "WORD_TABLE_CELL_LIMIT", "Word tables contain too many cells"
                    )
                lines.append("\t".join(values))
        core = document.core_properties
        return ParseResult(
            title=core.title or evidence.original_filename,
            text=ensure_text_limit("\n".join(lines)),
            language=core.language or None,
            metadata={
                "author": core.author or None,
                "subject": core.subject or None,
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "table_cell_count": table_cells,
            },
        )
