import asyncio
import hashlib
import uuid
from email.message import EmailMessage
from io import BytesIO

import pytest
from docx import Document  # type: ignore[import-untyped]
from httpx import AsyncClient
from openpyxl import Workbook  # type: ignore[import-untyped]
from pypdf import PdfWriter


@pytest.mark.integration
async def test_full_evidence_and_audit_flow(api_client: AsyncClient) -> None:
    created = await api_client.post("/api/v1/cases", json={"name": f"证据链测试-{uuid.uuid4()}"})
    assert created.status_code == 201, created.text
    case_id = created.json()["id"]
    content = (
        b"From: fictional@example.com\r\nTo: investigator@example.test\r\n"
        b"Subject: fictional shipment\r\n\r\nRelease No: TEST-0001"
    )
    files = {"file": ("sample.eml", content, "message/rfc822")}
    uploaded = await api_client.post(
        f"/api/v1/cases/{case_id}/evidence",
        files=files,
        data={"source_type": "EMAIL"},
    )
    assert uploaded.status_code == 201, uploaded.text
    evidence = uploaded.json()
    assert evidence["sha256"] == hashlib.sha256(content).hexdigest()
    assert evidence["object_key"].startswith(f"cases/{case_id}/evidence/")

    fetched = await api_client.get(f"/api/v1/cases/{case_id}/evidence/{evidence['id']}")
    assert fetched.status_code == 200
    assert fetched.json()["original_filename"] == "sample.eml"

    duplicate = await api_client.post(
        f"/api/v1/cases/{case_id}/evidence",
        files={"file": ("copy.eml", content, "message/rfc822")},
        data={"source_type": "EMAIL"},
    )
    assert duplicate.status_code == 409
    assert duplicate.json()["error"]["code"] == "DUPLICATE_EVIDENCE"

    audits = await api_client.get(f"/api/v1/cases/{case_id}/audit")
    assert audits.status_code == 200
    event_types = {event["event_type"] for event in audits.json()}
    assert {
        "CASE_CREATED",
        "EVIDENCE_UPLOAD_STARTED",
        "EVIDENCE_HASHED",
        "EVIDENCE_STORED",
        "EVIDENCE_CREATED",
    } <= event_types
    verification = await api_client.get(f"/api/v1/cases/{case_id}/audit/verify")
    assert verification.status_code == 200
    assert verification.json()["valid"] is True


@pytest.mark.integration
async def test_eml_processing_worker_normalizes_and_extracts_attachment(
    api_client: AsyncClient,
) -> None:
    created = await api_client.post(
        "/api/v1/cases", json={"name": f"邮件标准化测试-{uuid.uuid4()}"}
    )
    assert created.status_code == 201, created.text
    case_id = created.json()["id"]
    message = EmailMessage()
    message["From"] = "sender@example.test"
    message["To"] = "analyst@example.test"
    message["Subject"] = "标准化验收邮件"
    message["Message-ID"] = f"<{uuid.uuid4()}@example.test>"
    message.set_content("用于验证 MinIO 到 Worker 再到标准化文档的正文。")
    message.add_attachment(
        b"unique attachment content " + uuid.uuid4().bytes,
        maintype="text",
        subtype="plain",
        filename="clue.txt",
    )
    uploaded = await api_client.post(
        f"/api/v1/cases/{case_id}/evidence",
        files={"file": ("worker-sample.eml", message.as_bytes(), "message/rfc822")},
        data={"source_type": "EMAIL"},
    )
    assert uploaded.status_code == 201, uploaded.text
    evidence_id = uploaded.json()["id"]
    enqueued = await api_client.post(
        f"/api/v1/cases/{case_id}/evidence/{evidence_id}/processing-jobs",
        json={"idempotency_key": str(uuid.uuid4()), "max_attempts": 3},
    )
    assert enqueued.status_code == 202, enqueued.text

    job = enqueued.json()
    for _ in range(30):
        jobs = await api_client.get(
            f"/api/v1/cases/{case_id}/evidence/processing-jobs",
            params={"evidence_id": evidence_id},
        )
        assert jobs.status_code == 200, jobs.text
        job = jobs.json()[0]
        if job["status"] not in {"QUEUED", "PROCESSING"}:
            break
        await asyncio.sleep(0.25)

    assert job["status"] == "COMPLETED", job
    normalized = await api_client.get(f"/api/v1/cases/{case_id}/evidence/{evidence_id}/normalized")
    assert normalized.status_code == 200, normalized.text
    document = normalized.json()
    assert document["status"] == "READY"
    assert document["title"] == "标准化验收邮件"
    assert "MinIO" in document["text_preview"]
    assert document["metadata_json"]["attachment_count"] == 1
    assert document["content_object_key"].endswith(".json")

    evidence_list = await api_client.get(f"/api/v1/cases/{case_id}/evidence")
    assert evidence_list.status_code == 200
    children = [item for item in evidence_list.json() if item["parent_evidence_id"] == evidence_id]
    assert len(children) == 1
    assert children[0]["original_filename"] == "clue.txt"


def _office_samples() -> tuple[tuple[str, bytes, str, str], ...]:
    workbook_stream = BytesIO()
    workbook = Workbook()
    workbook.active.append(("账号", "金额"))
    workbook.active.append(("A-100", 500))
    workbook.save(workbook_stream)
    workbook.close()

    document_stream = BytesIO()
    document = Document()
    document.add_paragraph("DOCX 研判内容")
    document.save(document_stream)

    pdf_stream = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.write(pdf_stream)
    return (
        ("note.txt", "TXT 研判内容".encode(), "text/plain", "builtin-text"),
        ("records.csv", "姓名,金额\n张三,100\n".encode(), "text/csv", "builtin-csv"),
        (
            "ledger.xlsx",
            workbook_stream.getvalue(),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "python-excel",
        ),
        (
            "report.docx",
            document_stream.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "python-docx",
        ),
        ("scan.pdf", pdf_stream.getvalue(), "application/pdf", "pypdf"),
    )


@pytest.mark.integration
async def test_document_parsers_run_through_durable_worker(api_client: AsyncClient) -> None:
    created = await api_client.post(
        "/api/v1/cases", json={"name": f"多格式解析测试-{uuid.uuid4()}"}
    )
    assert created.status_code == 201, created.text
    case_id = created.json()["id"]
    expected_parsers: dict[str, str] = {}
    for filename, content, mime_type, parser_name in _office_samples():
        uploaded = await api_client.post(
            f"/api/v1/cases/{case_id}/evidence",
            files={"file": (filename, content, mime_type)},
            data={"source_type": "OTHER"},
        )
        assert uploaded.status_code == 201, uploaded.text
        evidence_id = uploaded.json()["id"]
        expected_parsers[evidence_id] = parser_name
        enqueued = await api_client.post(
            f"/api/v1/cases/{case_id}/evidence/{evidence_id}/processing-jobs",
            json={"idempotency_key": str(uuid.uuid4()), "max_attempts": 3},
        )
        assert enqueued.status_code == 202, enqueued.text

    jobs_by_evidence: dict[str, dict[str, object]] = {}
    for _ in range(80):
        response = await api_client.get(f"/api/v1/cases/{case_id}/evidence/processing-jobs")
        assert response.status_code == 200, response.text
        jobs_by_evidence = {item["evidence_id"]: item for item in response.json()}
        if len(jobs_by_evidence) == len(expected_parsers) and all(
            item["status"] not in {"QUEUED", "PROCESSING"} for item in jobs_by_evidence.values()
        ):
            break
        await asyncio.sleep(0.25)

    for evidence_id, parser_name in expected_parsers.items():
        job = jobs_by_evidence[evidence_id]
        assert job["status"] == "COMPLETED", job
        assert job["parser_name"] == parser_name
        normalized = await api_client.get(
            f"/api/v1/cases/{case_id}/evidence/{evidence_id}/normalized"
        )
        assert normalized.status_code == 200, normalized.text
        assert normalized.json()["status"] == "READY"


def _customs_case_samples() -> tuple[tuple[str, bytes, str], ...]:
    invoice_stream = BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(("发票号", "INV-2026-0821"))
    sheet.append(("报关单号", "CN202608210001"))
    sheet.append(("实际总价", 105000))
    workbook.save(invoice_stream)
    workbook.close()

    message = EmailMessage()
    message["From"] = "supplier@example.test"
    message["To"] = "purchase@example.test"
    message["Subject"] = "Final payment for INV-2026-0821"
    message.set_content(
        "合同实际总价为 105,000 美元。报关材料请按 68,000 美元准备申报。报关单号 CN202608210001。"
    )
    message.add_attachment(
        invoice_stream.getvalue(),
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename="INV-2026-0821.xlsx",
    )

    inspection_stream = BytesIO()
    inspection = Document()
    inspection.add_paragraph("报关单号 CN202608210001，实际数量与申报数量一致。")
    inspection.save(inspection_stream)
    return (
        ("03_供应商邮件.eml", message.as_bytes(), "message/rfc822"),
        (
            "04_付款记录.csv",
            (
                "交易序号,发票号,金额(USD)\n"
                "PAY-001,INV-2026-0821,68000\n"
                "PAY-002,INV-2026-0821,37000\n"
            ).encode(),
            "text/csv",
        ),
        (
            "05_查验记录.docx",
            inspection_stream.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
    )


@pytest.mark.integration
async def test_customs_case_runs_from_files_to_review_artifact(api_client: AsyncClient) -> None:
    created = await api_client.post(
        "/api/v1/cases", json={"name": f"海关价格风险端到端-{uuid.uuid4()}"}
    )
    assert created.status_code == 201, created.text
    case_id = created.json()["id"]
    for filename, content, mime_type in _customs_case_samples():
        uploaded = await api_client.post(
            f"/api/v1/cases/{case_id}/evidence",
            files={"file": (filename, content, mime_type)},
            data={"source_type": "CUSTOMS"},
        )
        assert uploaded.status_code == 201, uploaded.text
        evidence_id = uploaded.json()["id"]
        enqueued = await api_client.post(
            f"/api/v1/cases/{case_id}/evidence/{evidence_id}/processing-jobs",
            json={"idempotency_key": str(uuid.uuid4()), "max_attempts": 3},
        )
        assert enqueued.status_code == 202, enqueued.text

    jobs: list[dict[str, object]] = []
    for _ in range(100):
        response = await api_client.get(f"/api/v1/cases/{case_id}/evidence/processing-jobs")
        assert response.status_code == 200, response.text
        jobs = response.json()
        if len(jobs) == 4 and all(item["status"] not in {"QUEUED", "PROCESSING"} for item in jobs):
            break
        await asyncio.sleep(0.25)
    assert len(jobs) == 4
    assert all(item["status"] == "COMPLETED" for item in jobs), jobs

    started = await api_client.post(
        f"/api/v1/cases/{case_id}/workflows",
        json={"analysis_scope": "customs_risk_analysis"},
    )
    assert started.status_code == 201, started.text
    workflow = started.json()
    assert workflow["status"] == "WAITING_REVIEW"
    assert workflow["next_nodes"] == ["human_review"]
    analysis = workflow["state"]["customs_analysis"]
    assert analysis["risk_level"] == "HIGH"
    assert analysis["declared_amount_usd"] == 68000.0
    assert analysis["actual_amount_usd"] == 105000.0
    assert analysis["difference_usd"] == 37000.0
    assert len(analysis["evidence_refs"]) == 4

    timeline = await api_client.get(f"/api/v1/workflows/{workflow['thread_id']}/timeline")
    assert timeline.status_code == 200, timeline.text
    artifacts = timeline.json()["artifacts"]
    customs_artifacts = [
        item for item in artifacts if item["artifact_type"] == "CUSTOMS_RISK_ANALYSIS"
    ]
    assert len(customs_artifacts) == 1
    assert customs_artifacts[0]["content_json"]["difference_usd"] == 37000.0


@pytest.mark.integration
async def test_customs_workflow_rejects_start_until_evidence_is_ready(
    api_client: AsyncClient,
) -> None:
    created = await api_client.post("/api/v1/cases", json={"name": f"海关证据门禁-{uuid.uuid4()}"})
    case_id = created.json()["id"]
    uploaded = await api_client.post(
        f"/api/v1/cases/{case_id}/evidence",
        files={"file": ("待解析.txt", b"pending", "text/plain")},
        data={"source_type": "CUSTOMS"},
    )
    assert uploaded.status_code == 201, uploaded.text

    started = await api_client.post(
        f"/api/v1/cases/{case_id}/workflows",
        json={"analysis_scope": "customs_risk_analysis"},
    )

    assert started.status_code == 409, started.text
    assert started.json()["error"]["code"] == "EVIDENCE_NOT_READY"
