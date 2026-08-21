import csv
import uuid
from pathlib import Path

import pytest
from docx import Document  # type: ignore[import-untyped]
from openpyxl import Workbook  # type: ignore[import-untyped]
from pypdf import PdfWriter

from app.db.types import DocumentType, EvidenceSourceType
from app.models.evidence import Evidence
from app.parsers.csv_file import CsvParser
from app.parsers.pdf import PdfParser
from app.parsers.registry import build_parser_registry
from app.parsers.spreadsheet import SpreadsheetParser
from app.parsers.text import TextParser
from app.parsers.word import DocxParser


def evidence_for(source: Path, document_type: DocumentType) -> Evidence:
    return Evidence(
        id=uuid.uuid4(),
        case_id=uuid.uuid4(),
        original_filename=source.name,
        stored_filename=source.name,
        object_key=f"test/{source.name}",
        mime_type="application/octet-stream",
        file_extension=source.suffix.lower(),
        file_size=source.stat().st_size,
        sha256="0" * 64,
        source_type=EvidenceSourceType.OTHER,
        document_type=document_type,
        metadata_json={},
    )


def test_text_parser_supports_utf8_and_gb18030(tmp_path: Path) -> None:
    utf8 = tmp_path / "utf8.txt"
    utf8.write_text("中文 UTF-8", encoding="utf-8")
    legacy = tmp_path / "legacy.txt"
    legacy.write_bytes("中文 GB18030".encode("gb18030"))

    assert "中文 UTF-8" in TextParser().parse(utf8, evidence_for(utf8, DocumentType.TEXT)).text
    result = TextParser().parse(legacy, evidence_for(legacy, DocumentType.TEXT))
    assert "中文 GB18030" in result.text
    assert result.metadata["encoding"] == "gb18030"


def test_csv_parser_extracts_rows_and_shape(tmp_path: Path) -> None:
    source = tmp_path / "records.csv"
    with source.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerows((("姓名", "金额"), ("张三", "1200")))

    result = CsvParser().parse(source, evidence_for(source, DocumentType.CSV))

    assert "张三\t1200" in result.text
    assert result.metadata["row_count"] == 2
    assert result.metadata["column_count"] == 2


def test_xlsx_parser_extracts_multiple_sheets(tmp_path: Path) -> None:
    source = tmp_path / "ledger.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "交易"
    sheet.append(("账号", "金额"))
    sheet.append(("A-001", 88.5))
    workbook.create_sheet("说明").append(("测试",))
    workbook.save(source)
    workbook.close()

    result = SpreadsheetParser().parse(source, evidence_for(source, DocumentType.EXCEL))

    assert "## Sheet: 交易" in result.text
    assert "A-001\t88.5" in result.text
    assert result.metadata["sheet_count"] == 2


def test_docx_parser_extracts_paragraphs_and_tables(tmp_path: Path) -> None:
    source = tmp_path / "report.docx"
    document = Document()
    document.core_properties.title = "测试报告"
    document.add_paragraph("关键事实")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "对象"
    table.cell(0, 1).text = "张三"
    document.save(source)

    result = DocxParser().parse(source, evidence_for(source, DocumentType.WORD))

    assert result.title == "测试报告"
    assert "关键事实" in result.text
    assert "对象\t张三" in result.text
    assert result.metadata["table_count"] == 1


def test_pdf_parser_reads_page_metadata_and_recommends_ocr(tmp_path: Path) -> None:
    source = tmp_path / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.add_metadata({"/Title": "扫描件"})
    with source.open("wb") as handle:
        writer.write(handle)

    result = PdfParser().parse(source, evidence_for(source, DocumentType.PDF))

    assert result.title == "扫描件"
    assert result.metadata["page_count"] == 1
    assert result.metadata["ocr_recommended"] is True


@pytest.mark.parametrize(
    ("suffix", "expected_parser"),
    (
        (".txt", "builtin-text"),
        (".csv", "builtin-csv"),
        (".xlsx", "python-excel"),
        (".docx", "python-docx"),
        (".pdf", "pypdf"),
    ),
)
def test_registry_routes_supported_documents(
    tmp_path: Path, suffix: str, expected_parser: str
) -> None:
    source = tmp_path / f"sample{suffix}"
    source.write_bytes(b"placeholder")
    parser = build_parser_registry().find(evidence_for(source, DocumentType.UNKNOWN))
    assert parser is not None
    assert parser.name == expected_parser


def test_registry_leaves_legacy_doc_unsupported(tmp_path: Path) -> None:
    source = tmp_path / "legacy.doc"
    source.write_bytes(b"legacy")
    assert build_parser_registry().find(evidence_for(source, DocumentType.WORD)) is None
